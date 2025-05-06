import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter import ttk
import pandas as pd
from docx import Document
import os
import re
import logging
import datetime
import numpy as np
import openpyxl

# 设置日志记录
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(message)s')
logger = logging.getLogger(__name__)

# 格式化Excel数据
def format_cell_value(cell):
    """
    根据Excel单元格格式对数据进行格式化
    """
    if cell.value is None:
        return ""
    
    # 获取单元格格式
    format_code = cell.number_format
    value = cell.value
    
    # 处理文本格式
    if isinstance(value, str):
        return value
    
    # 处理日期时间类型
    if isinstance(value, (datetime.datetime, datetime.date)):
        if hasattr(value, 'hour') and (value.hour != 0 or value.minute != 0 or value.second != 0):
            return value.strftime('%Y-%m-%d %H:%M:%S')
        else:
            return value.strftime('%Y-%m-%d')
    
    # 处理数字类型
    if isinstance(value, (int, float)):
        # 处理前导零格式 (如 "000")
        if format_code.startswith('0') and not any(c != '0' for c in format_code):
            width = len(format_code)
            if isinstance(value, int) or value.is_integer():
                return str(int(value)).zfill(width)
        
        # 处理货币格式
        if any(s in format_code for s in ['¥', '$', '€', '￥']):
            decimal_places = 2  # 默认保留2位小数
            if '.' in format_code:
                # 尝试从格式中提取小数位数
                decimal_part = format_code.split('.')[-1]
                if '0' in decimal_part:
                    decimal_places = decimal_part.count('0')
            
            # 获取货币符号
            symbol = '¥'  # 默认人民币符号
            for s in ['¥', '$', '€', '￥']:
                if s in format_code:
                    symbol = s
                    break
            
            if value.is_integer() and decimal_places == 0:
                return f"{symbol}{int(value):,}"
            else:
                return f"{symbol}{value:,.{decimal_places}f}"
        
        # 处理带千分位的数字
        if '#,##0' in format_code or '#,###' in format_code:
            decimal_places = 0
            if '.' in format_code:
                decimal_part = format_code.split('.')[-1]
                if '0' in decimal_part:
                    decimal_places = decimal_part.count('0')
            
            if value.is_integer() and decimal_places == 0:
                return f"{int(value):,}"
            else:
                return f"{value:,.{decimal_places}f}"
        
        # 处理百分比
        if '%' in format_code:
            decimal_places = 0
            if '.' in format_code:
                decimal_part = format_code.split('.')[-1]
                if '0' in decimal_part:
                    decimal_places = decimal_part.count('0')
            return f"{value*100:.{decimal_places}f}%"
        
        # 处理普通数字
        if value.is_integer():
            return str(int(value))
        
        # 移除不必要的小数点和尾随零
        return str(value).rstrip('0').rstrip('.') if '.' in str(value) else str(value)
    
    # 其他情况，转换为字符串
    return str(value)

# 读取Excel数据和格式信息
def read_excel_with_format(excel_path):
    """
    同时读取Excel数据和单元格格式信息
    """
    # 读取数据
    df = pd.read_excel(excel_path)
    
    # 读取格式信息
    wb = openpyxl.load_workbook(excel_path, data_only=True)
    ws = wb.active
    
    # 获取列名
    column_names = list(df.columns)
    
    # 存储格式化后的数据
    formatted_data = []
    
    # 遍历行
    for row_idx in range(len(df)):
        # Excel行索引从1开始，且第1行是标题
        xl_row = row_idx + 2
        row_data = {}
        
        # 遍历列
        for col_idx, col_name in enumerate(column_names):
            # Excel列索引从1开始
            xl_col = col_idx + 1
            cell = ws.cell(row=xl_row, column=xl_col)
            
            # 格式化单元格值
            formatted_value = format_cell_value(cell)
            row_data[col_name] = formatted_value
        
        formatted_data.append(row_data)
    
    wb.close()
    return df, formatted_data

# 提取Word文档中的所有占位符
def extract_placeholders(doc):
    placeholder_pattern = re.compile(r'«([^»]+)»')
    placeholders = set()
    
    # 从段落中提取
    for para in doc.paragraphs:
        para_text = ''.join([run.text for run in para.runs])
        for match in placeholder_pattern.finditer(para_text):
            placeholders.add(match.group(1))
    
    # 从表格中提取
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para_text = ''.join([run.text for run in para.runs])
                    for match in placeholder_pattern.finditer(para_text):
                        placeholders.add(match.group(1))
    
    # 从页眉页脚中提取
    for section in doc.sections:
        for part in [section.header, section.footer]:
            # 从段落中提取
            for para in part.paragraphs:
                para_text = ''.join([run.text for run in para.runs])
                for match in placeholder_pattern.finditer(para_text):
                    placeholders.add(match.group(1))
            
            # 从表格中提取
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            para_text = ''.join([run.text for run in para.runs])
                            for match in placeholder_pattern.finditer(para_text):
                                placeholders.add(match.group(1))
    
    return placeholders

# 替换模板中的字段
def replace_placeholders(doc, replacements):
    # 定义正则表达式模式，用于匹配占位符
    placeholder_pattern = re.compile(r'«([^»]+)»')
    replaced_count = 0
    
    try:
        # 替换段落中的占位符（更好地保留格式）
        for para_idx, para in enumerate(doc.paragraphs):
            # 收集所有runs的文本和长度信息
            run_texts = [run.text for run in para.runs]
            para_text = ''.join(run_texts)
            
            # 查找段落中的所有占位符
            matches = list(placeholder_pattern.finditer(para_text))
            
            # 如果段落中没有占位符，则跳过
            if not matches:
                continue
                
            # 从后向前替换，避免位置偏移问题
            for match in reversed(matches):
                placeholder = match.group(0)  # 完整的占位符，如 «name»
                key = match.group(1)          # 占位符中的键，如 name
                
                if key in replacements:
                    value = replacements[key]
                    
                    # 确定占位符在文本中的位置
                    start_pos = match.start()
                    end_pos = match.end()
                    
                    # 确定占位符跨越的runs
                    start_run_idx = None
                    end_run_idx = None
                    current_pos = 0
                    
                    # 确定占位符开始和结束的run索引
                    for i, run_text in enumerate(run_texts):
                        run_len = len(run_text)
                        if start_run_idx is None and current_pos <= start_pos < current_pos + run_len:
                            start_run_idx = i
                            start_run_pos = start_pos - current_pos
                        
                        if end_run_idx is None and current_pos < end_pos <= current_pos + run_len:
                            end_run_idx = i
                            end_run_pos = end_pos - current_pos
                            break
                        
                        current_pos += run_len
                    
                    # 如果找到了占位符的位置
                    if start_run_idx is not None and end_run_idx is not None:
                        try:
                            # 处理占位符在单个run中的情况
                            if start_run_idx == end_run_idx:
                                run = para.runs[start_run_idx]
                                run.text = run.text[:start_run_pos] + value + run.text[end_run_pos:]
                            else:
                                # 处理占位符跨越多个runs的情况
                                # 处理第一个run（保留前部分并添加替换值）
                                first_run = para.runs[start_run_idx]
                                first_run.text = first_run.text[:start_run_pos] + value
                                
                                # 处理最后一个run（只保留尾部）
                                last_run = para.runs[end_run_idx]
                                last_run.text = last_run.text[end_run_pos:]
                                
                                # 清空中间的runs
                                for i in range(start_run_idx + 1, end_run_idx):
                                    para.runs[i].text = ""
                            
                            replaced_count += 1
                        except Exception as e:
                            logger.error(f"替换段落#{para_idx}中的'{placeholder}'时出错: {e}")
        
        # 替换表格中的占位符
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        # 收集所有runs的文本和长度信息
                        if len(paragraph.runs) == 0:
                            continue
                            
                        run_texts = [run.text for run in paragraph.runs]
                        para_text = ''.join(run_texts)
                        
                        # 查找段落中的所有占位符
                        matches = list(placeholder_pattern.finditer(para_text))
                        
                        # 如果段落中没有占位符，则跳过
                        if not matches:
                            continue
                            
                        # 从后向前替换，避免位置偏移问题
                        for match in reversed(matches):
                            placeholder = match.group(0)  # 完整的占位符，如 «name»
                            key = match.group(1)          # 占位符中的键，如 name
                            
                            if key in replacements:
                                value = replacements[key]
                                
                                # 确定占位符在文本中的位置
                                start_pos = match.start()
                                end_pos = match.end()
                                
                                # 确定占位符跨越的runs
                                start_run_idx = None
                                end_run_idx = None
                                current_pos = 0
                                
                                # 确定占位符开始和结束的run索引
                                for i, run_text in enumerate(run_texts):
                                    run_len = len(run_text)
                                    if start_run_idx is None and current_pos <= start_pos < current_pos + run_len:
                                        start_run_idx = i
                                        start_run_pos = start_pos - current_pos
                                    
                                    if end_run_idx is None and current_pos < end_pos <= current_pos + run_len:
                                        end_run_idx = i
                                        end_run_pos = end_pos - current_pos
                                        break
                                    
                                    current_pos += run_len
                                
                                # 如果找到了占位符的位置
                                if start_run_idx is not None and end_run_idx is not None:
                                    try:
                                        # 处理占位符在单个run中的情况
                                        if start_run_idx == end_run_idx:
                                            run = paragraph.runs[start_run_idx]
                                            run.text = run.text[:start_run_pos] + value + run.text[end_run_pos:]
                                        else:
                                            # 处理占位符跨越多个runs的情况
                                            # 处理第一个run（保留前部分并添加替换值）
                                            first_run = paragraph.runs[start_run_idx]
                                            first_run.text = first_run.text[:start_run_pos] + value
                                            
                                            # 处理最后一个run（只保留尾部）
                                            last_run = paragraph.runs[end_run_idx]
                                            last_run.text = last_run.text[end_run_pos:]
                                            
                                            # 清空中间的runs
                                            for i in range(start_run_idx + 1, end_run_idx):
                                                paragraph.runs[i].text = ""
                                        
                                        replaced_count += 1
                                    except Exception as e:
                                        logger.error(f"替换表格#{table_idx}的单元格[{row_idx},{cell_idx}]段落#{para_idx}中的'{placeholder}'时出错: {e}")
        
        # 替换页眉和页脚中的占位符
        for section_idx, section in enumerate(doc.sections):
            # 处理页眉和页脚
            try:
                for part_name, part in [('header', section.header), ('footer', section.footer)]:
                    for para_idx, para in enumerate(part.paragraphs):
                        # 收集所有runs的文本
                        if len(para.runs) == 0:
                            continue
                            
                        run_texts = [run.text for run in para.runs]
                        para_text = ''.join(run_texts)
                        
                        # 查找段落中的所有占位符
                        matches = list(placeholder_pattern.finditer(para_text))
                        
                        # 如果段落中没有占位符，则跳过
                        if not matches:
                            continue
                            
                        # 从后向前替换，避免位置偏移问题
                        for match in reversed(matches):
                            placeholder = match.group(0)  # 完整的占位符，如 «name»
                            key = match.group(1)          # 占位符中的键，如 name
                            
                            if key in replacements:
                                value = replacements[key]
                                
                                # 确定占位符在文本中的位置
                                start_pos = match.start()
                                end_pos = match.end()
                                
                                # 确定占位符跨越的runs
                                start_run_idx = None
                                end_run_idx = None
                                current_pos = 0
                                
                                # 确定占位符开始和结束的run索引
                                for i, run_text in enumerate(run_texts):
                                    run_len = len(run_text)
                                    if start_run_idx is None and current_pos <= start_pos < current_pos + run_len:
                                        start_run_idx = i
                                        start_run_pos = start_pos - current_pos
                                    
                                    if end_run_idx is None and current_pos < end_pos <= current_pos + run_len:
                                        end_run_idx = i
                                        end_run_pos = end_pos - current_pos
                                        break
                                    
                                    current_pos += run_len
                                
                                # 如果找到了占位符的位置
                                if start_run_idx is not None and end_run_idx is not None:
                                    try:
                                        # 处理占位符在单个run中的情况
                                        if start_run_idx == end_run_idx:
                                            run = para.runs[start_run_idx]
                                            run.text = run.text[:start_run_pos] + value + run.text[end_run_pos:]
                                        else:
                                            # 处理占位符跨越多个runs的情况
                                            # 处理第一个run（保留前部分并添加替换值）
                                            first_run = para.runs[start_run_idx]
                                            first_run.text = first_run.text[:start_run_pos] + value
                                            
                                            # 处理最后一个run（只保留尾部）
                                            last_run = para.runs[end_run_idx]
                                            last_run.text = last_run.text[end_run_pos:]
                                            
                                            # 清空中间的runs
                                            for i in range(start_run_idx + 1, end_run_idx):
                                                para.runs[i].text = ""
                                        
                                        replaced_count += 1
                                    except Exception as e:
                                        logger.error(f"替换章节#{section_idx}的{part_name}段落#{para_idx}中的'{placeholder}'时出错: {e}")
            except Exception as e:
                logger.error(f"处理章节#{section_idx}的页眉或页脚时出错: {e}")
                        
        return replaced_count
    except Exception as e:
        logger.error(f"替换文档模板时发生错误: {e}")
        return 0

# 主界面类
class MailMergeApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Excel-Word 邮件合并工具")
        self.root.geometry("700x650")

        self.excel_path = ""
        self.word_path = ""
        self.output_dir = ""
        self.df = None
        self.formatted_data = None
        self.template_placeholders = set()  # 存储模板中的占位符

        # Excel 文件选择
        tk.Label(root, text="① 请选择 Excel 文件：").pack(pady=5)
        tk.Button(root, text="选择 Excel", command=self.select_excel).pack()

        # Word 模板选择
        tk.Label(root, text="② 请选择 Word 模板文件：").pack(pady=5)
        tk.Button(root, text="选择 Word 模板", command=self.select_word).pack()

        # 输出文件命名列选择
        tk.Label(root, text="③ 选择用于命名 Word 文件的列名：").pack(pady=5)
        self.filename_column = ttk.Combobox(root, state="readonly")
        self.filename_column.pack()

        # 输出路径选择
        tk.Label(root, text="④ 选择输出文件夹（可选）：").pack(pady=5)
        tk.Button(root, text="选择输出文件夹", command=self.select_output_dir).pack()
        self.output_dir_label = tk.Label(root, text="默认使用 Excel 同目录的 output_docs 文件夹")
        self.output_dir_label.pack()

        # 字段映射检查按钮
        tk.Button(root, text="⤷ 检查字段映射", command=self.check_field_mapping).pack(pady=5)

        # 合并执行按钮
        tk.Button(root, text="⑤ 开始合并生成文档", command=self.generate_docs, bg="green", fg="white").pack(pady=20)

        # 状态输出框
        self.status = tk.Text(root, height=10, width=120)
        self.status.pack()

    def select_excel(self):
        self.excel_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx *.xls")])
        if self.excel_path:
            try:
                # 读取Excel数据和格式信息
                self.df, self.formatted_data = read_excel_with_format(self.excel_path)
                self.status.insert(tk.END, f"✅ 已加载 Excel 文件：{self.excel_path}\n")
                self.status.insert(tk.END, f"   共 {len(self.df.columns)} 列, {len(self.df)} 行数据\n")
                self.filename_column['values'] = list(self.df.columns)
                if len(self.df.columns) > 0:
                    self.filename_column.current(0)
                
                # 如果已经选择了Word模板，检查字段映射
                if hasattr(self, 'template_placeholders') and self.template_placeholders:
                    self.check_field_mapping()
            except Exception as e:
                messagebox.showerror("错误", f"无法读取 Excel 文件：{e}")
                logger.error(f"读取Excel文件出错: {e}")

    def select_word(self):
        self.word_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
        if self.word_path:
            try:
                doc = Document(self.word_path)
                self.template_placeholders = extract_placeholders(doc)
                placeholder_count = len(self.template_placeholders)
                self.status.insert(tk.END, f"✅ 已选择 Word 模板文件：{self.word_path}，包含 {placeholder_count} 个不同占位符\n")
                
                # 如果已经选择了Excel文件，检查字段映射
                if self.df is not None:
                    self.check_field_mapping()
            except Exception as e:
                messagebox.showerror("错误", f"无法读取 Word 模板：{e}")

    def select_output_dir(self):
        self.output_dir = filedialog.askdirectory()
        if self.output_dir:
            self.output_dir_label.config(text=f"📁 输出目录：{self.output_dir}")

    def check_field_mapping(self):
        if not hasattr(self, 'template_placeholders') or not self.template_placeholders:
            if self.word_path:
                try:
                    doc = Document(self.word_path)
                    self.template_placeholders = extract_placeholders(doc)
                except Exception as e:
                    messagebox.showerror("错误", f"无法读取 Word 模板：{e}")
                    return
            else:
                messagebox.showwarning("警告", "请先选择 Word 模板文件！")
                return
        
        if self.df is None:
            messagebox.showwarning("警告", "请先选择 Excel 文件！")
            return
        
        # 检查字段映射
        excel_columns = set(self.df.columns)
        missing_fields = [field for field in self.template_placeholders if field not in excel_columns]
        
        self.status.insert(tk.END, "\n=== 字段映射检查 ===\n")
        
        if missing_fields:
            self.status.insert(tk.END, f"⚠️ 发现 {len(missing_fields)} 个占位符在 Excel 中没有对应列：\n")
            for field in missing_fields:
                self.status.insert(tk.END, f"  - «{field}»\n")
            
            messagebox.showwarning("字段映射警告", 
                                  f"发现 {len(missing_fields)} 个模板占位符在 Excel 中没有对应列。\n"
                                  f"这些占位符在生成文档时将保持不变。\n"
                                  f"详情请查看状态窗口。")
        else:
            self.status.insert(tk.END, "✅ 所有占位符在 Excel 中都有对应列！\n")
            messagebox.showinfo("字段映射正确", "所有占位符在 Excel 中都有对应列！")
        
        unused_columns = [col for col in excel_columns if col not in self.template_placeholders]
        if unused_columns:
            self.status.insert(tk.END, f"ℹ️ Excel 中有 {len(unused_columns)} 个列在模板中未使用\n")

    def generate_docs(self):
        if self.df is None or self.formatted_data is None or not self.word_path:
            messagebox.showwarning("警告", "请确保已选择 Excel 和 Word 模板！")
            return

        selected_column = self.filename_column.get()
        if not selected_column:
            messagebox.showwarning("警告", "请先选择用于命名文档的 Excel 列名！")
            return

        try:
            # 使用用户选择的输出路径或默认路径
            output_dir = self.output_dir or os.path.join(os.path.dirname(self.excel_path), "output_docs")
            os.makedirs(output_dir, exist_ok=True)
            
            successful_docs = 0
            issues = []
            
            # 显示进度信息
            self.status.insert(tk.END, f"⏳ 开始处理 {len(self.df)} 份文档...\n")
            self.root.update()

            # 获取Excel工作簿以读取原始格式
            wb = openpyxl.load_workbook(self.excel_path, data_only=True)
            ws = wb.active
            
            # 获取命名列的索引
            column_names = list(self.df.columns)
            filename_col_idx = column_names.index(selected_column)

            for i, row_data in enumerate(self.formatted_data):
                try:
                    doc = Document(self.word_path)
                    replaced = replace_placeholders(doc, row_data)

                    # 获取用于命名的单元格（Excel行从2开始，因为第1行是标题）
                    excel_row = i + 2
                    excel_col = filename_col_idx + 1  # Excel列从1开始
                    naming_cell = ws.cell(row=excel_row, column=excel_col)
                    
                    # 使用格式化后的单元格值作为文件名，保持格式
                    name_formatted = format_cell_value(naming_cell)
                    
                    # 确保文件名安全
                    name_str = str(name_formatted).strip().replace("/", "_").replace("\\", "_").replace(":", "_").replace("*", "_").replace("?", "_").replace("\"", "_").replace("<", "_").replace(">", "_").replace("|", "_")
                    
                    # 如果文件名为空，使用默认名称
                    if not name_str or name_str.isspace():
                        name_str = f"document_{i+1}"
                    
                    output_path = os.path.join(output_dir, f"{name_str}.docx")
                    doc.save(output_path)
                    
                    successful_docs += 1
                    
                    # 每10个文档更新一次状态
                    if i % 10 == 0 or i == len(self.df) - 1:
                        self.status.insert(tk.END, f"✓ 已生成 {i+1}/{len(self.df)} 份文档\n")
                        self.root.update()
                        
                except Exception as e:
                    error_msg = f"处理第 {i+1} 行数据时出错: {e}"
                    logger.error(error_msg)
                    issues.append(error_msg)
            
            # 关闭工作簿
            wb.close()

            self.status.insert(tk.END, f"🎉 成功生成 {successful_docs} 份文档，保存在：{output_dir}\n")
            
            if issues:
                self.status.insert(tk.END, f"⚠️ 处理过程中有 {len(issues)} 个问题\n")
                messagebox.showinfo("完成", f"成功生成 {successful_docs} 份文档，有 {len(issues)} 个问题。请查看日志了解详情。")
            else:
                messagebox.showinfo("完成", f"成功生成 {successful_docs} 份文档！")
        except Exception as e:
            error_msg = f"生成文档过程中发生错误: {e}"
            self.status.insert(tk.END, f"❌ {error_msg}\n")
            logger.error(error_msg)
            messagebox.showerror("错误", str(e))


# 启动界面
if __name__ == "__main__":
    root = tk.Tk()
    app = MailMergeApp(root)
    root.mainloop()
